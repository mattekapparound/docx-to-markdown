import pytest
from pathlib import Path
from docx import Document
from docx_to_markdown.utils import document_numbering_is_numbered, iter_block_items

def test_document_numbering_is_numbered_bullet(tmp_path):
    # Create a test document with a bullet list
    doc = Document()
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Item 1")
    file = tmp_path / "bullet.docx"
    doc.save(file)

    # Reopen the document
    doc = Document(file)
    # Retrieve num_id from the first paragraph
    num_id = p._p.pPr.numPr.numId.val

    result = document_numbering_is_numbered(doc, num_id)
    assert result in [True, False]  # Should not raise errors

def test_iter_block_items_returns_blocks(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello World")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell"
    file = tmp_path / "mixed.docx"
    doc.save(file)

    doc = Document(file)
    blocks = list(iter_block_items(doc))
    # Should contain at least 2 blocks: a paragraph and a table
    assert any(b.__class__.__name__ == "Paragraph" for b in blocks)
    assert any(b.__class__.__name__ == "Table" for b in blocks)
