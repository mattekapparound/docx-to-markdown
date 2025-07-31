from pathlib import Path
from docx import Document
from docx_to_markdown.converter import word_to_markdown

def test_word_to_markdown_headings(tmp_path):
    file = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Title 1", level=1)
    doc.add_heading("Subtitle", level=2)
    doc.save(file)

    md = word_to_markdown(str(file))
    assert "# Title 1" in md
    assert "## Subtitle" in md

def test_word_to_markdown_list(tmp_path):
    file = tmp_path / "list.docx"
    doc = Document()
    doc.add_paragraph("Item 1", style='List Bullet')
    doc.add_paragraph("Item 2", style='List Number')
    doc.save(file)

    md = word_to_markdown(str(file))
    assert "- Item 1" in md or "1. Item 1" in md
    assert "- Item 2" in md or "1. Item 2" in md

def test_word_to_markdown_table(tmp_path):
    file = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "V1"
    table.cell(1, 1).text = "V2"
    doc.save(file)

    md = word_to_markdown(str(file))
    assert "| H1 | H2 |" in md
    assert "| V1 | V2 |" in md
